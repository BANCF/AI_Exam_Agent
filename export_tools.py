import io
import zipfile
import json
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from utils import clean_latex_for_word

def create_word_document(exam_data_list):
    """Tạo file Word ĐỀ THI CHÍNH THỨC (Không chứa lời giải)"""
    doc = Document()
    p_header = doc.add_paragraph()
    p_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_so = p_header.add_run('SỞ GIÁO DỤC VÀ ĐÀO TẠO HÀ NỘI\n')
    run_so.font.name = 'Times New Roman'; run_so.font.size = Pt(13)
    run_ky_thi = p_header.add_run('KỲ THI CHỌN HỌC SINH GIỎI CẤP THÀNH PHỐ\n')
    run_ky_thi.font.name = 'Times New Roman'; run_ky_thi.font.size = Pt(13)
    run_nam = p_header.add_run('NĂM HỌC 2025 - 2026\n')
    run_nam.font.name = 'Times New Roman'; run_nam.font.size = Pt(12)
    doc.add_paragraph('________________________').alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p_mon = doc.add_paragraph()
    p_mon.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_mon = p_mon.add_run('\nĐỀ THI MÔN: TIN HỌC (LẬP TRÌNH)\n')
    run_mon.bold = True; run_mon.font.name = 'Times New Roman'; run_mon.font.size = Pt(16)
    run_tg = p_mon.add_run('Thời gian làm bài: 150 phút\n(Đề thi gồm có 0{} trang)\n'.format(len(exam_data_list)))
    run_tg.italic = True; run_tg.font.name = 'Times New Roman'; run_tg.font.size = Pt(12)

    doc.add_heading('TỔNG QUAN BÀI THI', level=2)
    table_tong_quan = doc.add_table(rows=1, cols=4)
    table_tong_quan.style = 'Table Grid'
    hdr = table_tong_quan.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = 'Tên bài', 'File chương trình', 'File dữ liệu vào', 'File dữ liệu ra'
    for c in hdr:
        for p in c.paragraphs:
            for r in p.runs: r.bold = True

    for idx, prob in enumerate(exam_data_list):
        row = table_tong_quan.add_row().cells
        file_name = prob.get('title', f'BAI{idx+1}').split()[0].upper()
        row[0].text = f"Bài {idx + 1}: {prob.get('title', '')}"
        row[1].text = f"{file_name}.CPP / .PY"
        row[2].text = f"{file_name}.INP"
        row[3].text = f"{file_name}.OUT"
        for i in range(1, 4):
            for p in row[i].paragraphs:
                for r in p.runs: r.font.name = 'Courier New'

    doc.add_paragraph("\n")

    for idx, prob in enumerate(exam_data_list):
        doc.add_heading(f"BÀI {idx + 1} ( {prob.get('title', 'Bài toán')} )", level=1)
        doc.add_paragraph(clean_latex_for_word(prob.get('legend', '')))
        doc.add_heading('Dữ liệu vào (Input):', level=3)
        doc.add_paragraph(clean_latex_for_word(prob.get('input_format', '')))
        doc.add_heading('Kết quả ra (Output):', level=3)
        doc.add_paragraph(clean_latex_for_word(prob.get('output_format', '')))
        doc.add_heading('Giới hạn:', level=3)
        doc.add_paragraph(clean_latex_for_word(prob.get('constraints', '')))
        
        doc.add_heading('Ví dụ (Sample Test Cases):', level=3)
        # Chỉ lấy tối đa 2 testcase đầu tiên làm ví dụ minh họa trong đề thi
        sample_tests = prob.get("test_cases", [])[:2]
        for i, tc in enumerate(sample_tests):
            doc.add_paragraph(f"Ví dụ {i+1}:", style='List Bullet')
            table = doc.add_table(rows=2, cols=2)
            table.style = 'Table Grid'
            table.rows[0].cells[0].text = 'Input'; table.rows[0].cells[1].text = 'Output'
            table.rows[1].cells[0].paragraphs[0].add_run(tc.get('input_data', '')).font.name = 'Courier New'
            table.rows[1].cells[1].paragraphs[0].add_run(tc.get('output_data', '')).font.name = 'Courier New'
            doc.add_paragraph(f"Giải thích: {clean_latex_for_word(tc.get('explanation', ''))}\n")
        
        if idx < len(exam_data_list) - 1: doc.add_page_break()

    byte_io = io.BytesIO()
    doc.save(byte_io)
    byte_io.seek(0)
    return byte_io.getvalue()


def create_editorial_document(exam_data_list):
    """Tạo file Word ĐÁP ÁN & HƯỚNG DẪN CHẤM CHI TIẾT (Dành riêng cho giáo viên)"""
    doc = Document()
    p_header = doc.add_paragraph()
    p_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_so = p_header.add_run('SỞ GIÁO DỤC VÀ ĐÀO TẠO HÀ NỘI\n')
    run_so.font.name = 'Times New Roman'; run_so.font.size = Pt(13)
    run_title = p_header.add_run('HƯỚNG DẪN CHẤM - ĐÁP ÁN KỲ THI HỌC SINH GIỎI\n')
    run_title.bold = True; run_title.font.name = 'Times New Roman'; run_title.font.size = Pt(15)
    doc.add_paragraph('________________________').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("\n")

    for idx, prob in enumerate(exam_data_list):
        doc.add_heading(f"BÀI {idx + 1}: {prob.get('title', 'Bài toán')}", level=1)
        
        # 1. Hướng dẫn thuật toán (Editorial)
        doc.add_heading('1. Ý tưởng giải thuật & Đánh giá độ phức tạp:', level=3)
        doc.add_paragraph(clean_latex_for_word(prob.get('editorial', 'Chưa cập nhật hướng dẫn giải thuật.')))
        
        # 2. Gợi ý làm bài (Hints)
        doc.add_heading('2. Các bước gợi ý tư duy (Hints):', level=3)
        hints = prob.get('hints', [])
        if hints:
            for i, hint in enumerate(hints):
                doc.add_paragraph(f"- Gợi ý {i+1}: {clean_latex_for_word(hint)}")
        else:
            doc.add_paragraph("Không có gợi ý bổ sung.")

        # 3. Đáp án Code C++ mẫu
        doc.add_heading('3. Chương trình giải mẫu bằng C++:', level=3)
        cpp_code = prob.get('c_plus_plus_solution', '')
        if cpp_code:
            p_cpp = doc.add_paragraph()
            r_cpp = p_cpp.add_run(cpp_code)
            r_cpp.font.name = 'Courier New'; r_cpp.font.size = Pt(10)
        else:
            doc.add_paragraph("Chưa cấu hình code C++ mẫu.")

        # 4. Đáp án Code Python mẫu
        doc.add_heading('4. Chương trình giải mẫu bằng Python:', level=3)
        py_code = prob.get('python_solution', '')
        if py_code:
            p_py = doc.add_paragraph()
            r_py = p_py.add_run(py_code)
            r_py.font.name = 'Courier New'; r_py.font.size = Pt(10)
        else:
            doc.add_paragraph("Chưa cấu hình code Python mẫu.")

        if idx < len(exam_data_list) - 1: doc.add_page_break()

    byte_io = io.BytesIO()
    doc.save(byte_io)
    byte_io.seek(0)
    return byte_io.getvalue()


def create_themis_zip(exam_cart):
    """
    Tự động đóng gói toàn bộ danh sách câu hỏi trong đề thành:
    - File Đề bài Word sạch.
    - File Hướng dẫn giải + Đáp án Word.
    - Cấu trúc thư mục Test01 -> Test10 chuẩn phần mềm chấm Themis.
    """
    zip_buffer = io.BytesIO()
    
    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
        
        # --- BƯỚC THÊM MỚI: Sinh tài liệu văn bản lưu vào gốc File ZIP ---
        if exam_cart:
            # Sinh file Đề bài Word không đáp án
            de_bai_bytes = create_word_document(exam_cart)
            zip_file.writestr("De_Bai_Chinh_Thuc.docx", de_bai_bytes)
            
            # Sinh file Đáp án + Lời giải Word độc lập
            dap_an_bytes = create_editorial_document(exam_cart)
            zip_file.writestr("Huong_Dan_Cham_Va_Dap_An.docx", dap_an_bytes)
        
        # --- BƯỚC CŨ: Đóng gói cấu trúc Thư mục chấm test tự động ---
        for idx, problem in enumerate(exam_cart):
            raw_title = problem.get("title", f"BAI{idx+1}")
            folder_name = raw_title.split()[0].replace(":", "").upper()
            base_path = f"{folder_name}/"
            
            cpp_solution = problem.get("c_plus_plus_solution", "")
            python_solution = problem.get("python_solution", "")
            
            if cpp_solution:
                zip_file.writestr(f"{base_path}{folder_name}.cpp", cpp_solution)
            if python_solution:
                zip_file.writestr(f"{base_path}{folder_name}.py", python_solution)
            if "generator_script" in problem:
                zip_file.writestr(f"{base_path}generator.py", problem["generator_script"])
                
            test_cases = problem.get("test_cases", [])
            for tc_idx, tc in enumerate(test_cases):
                test_folder_name = f"Test{list_index_to_str(tc_idx + 1)}"
                test_path = f"{base_path}{test_folder_name}/"
                
                input_data = tc.get("input_data", "").strip()
                output_data = tc.get("output_data", "").strip()
                
                zip_file.writestr(f"{test_path}{folder_name}.INP", input_data)
                zip_file.writestr(f"{test_path}{folder_name}.OUT", output_data)
                
    return zip_buffer.getvalue()

def list_index_to_str(index):
    """Hàm bổ trợ chuyển số 1 -> '01', 9 -> '09', 10 -> '10' để đặt tên folder chuẩn"""
    return f"0{index}" if index < 10 else str(index)